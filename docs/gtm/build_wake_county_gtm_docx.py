from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


OUTPUT_PATH = "docs/gtm/HomeCare-Wake-County-Offline-GTM-2026-04-06.docx"


TARGETS = {
    "Priority 1: Senior Centers And Core Referral Organizations": [
        {
            "name": "Five Points Center for Active Adults",
            "town": "Raleigh",
            "type": "Senior center",
            "first_move": "Call first Monday morning, then visit",
            "contact": "2000 Noble Rd, Raleigh, NC 27608 | 919-996-4730 | raleighnc.gov/parks-and-recreation/places/five-points-center-active-adults",
            "ask": "Flyer approval, newsletter inclusion, short talk",
            "notes": "High-trust active adult site. Meals on Wheels already has presence there.",
        },
        {
            "name": "Anne Gordon Center for Active Adults",
            "town": "Raleigh",
            "type": "Senior center",
            "first_move": "Call first Monday morning, then visit",
            "contact": "1901 Spring Forest Rd, Raleigh, NC 27615 | 919-996-4720 | Brian Philpot | raleighnc.gov/parks-and-recreation/places/anne-gordon-center-active-adults",
            "ask": "Flyer approval, program lead intro, educational session",
            "notes": "Resources for Seniors has appointments there. Strong trust-transfer location.",
        },
        {
            "name": "Cary Senior Center",
            "town": "Cary",
            "type": "Senior center",
            "first_move": "Call first Tuesday morning, then visit",
            "contact": "120 Maury O'Dell Place, Cary, NC 27513 | 919-469-4081 | carync.gov/recreation-enjoyment/facilities/senior-center",
            "ask": "Flyer approval, bulletin/newsletter, event-table intro",
            "notes": "High-priority Cary family audience.",
        },
        {
            "name": "Apex Senior Center",
            "town": "Apex",
            "type": "Senior center",
            "first_move": "Call first Tuesday morning, then visit",
            "contact": "63 Hunter St, Apex, NC 27502 | 919-249-3354 | apexnc.org/seniorcenter",
            "ask": "Flyer approval, event-table lead, educational session",
            "notes": "Official page invites program proposals.",
        },
        {
            "name": "Garner Senior Center",
            "town": "Garner",
            "type": "Senior center",
            "first_move": "Call first Wednesday morning, then visit",
            "contact": "205 E Garner Rd, Garner, NC 27529 | 919-779-0122 | garnerseniorcenter@garnernc.gov | garnernc.gov",
            "ask": "Flyer approval, newsletter, sponsor/event intro",
            "notes": "Resources for Seniors and Meals on Wheels already intersect here.",
        },
        {
            "name": "Wake Forest Center for Active Aging",
            "town": "Wake Forest",
            "type": "Senior center",
            "first_move": "Call first Wednesday morning, then visit",
            "contact": "235 E Holding Ave, Wake Forest, NC 27587 | 919-554-4111 | wakeforestnc.gov/wake-forest-center-active-aging",
            "ask": "Flyer approval, community partner conversation, outreach slot",
            "notes": "Official site has a clear community-partner path.",
        },
        {
            "name": "Resources for Seniors",
            "town": "Raleigh",
            "type": "Referral nonprofit",
            "first_move": "Call first Monday morning, request meeting, then visit",
            "contact": "1110 Navaho Dr, Fourth Floor, Raleigh, NC 27609 | 919-872-7933 | info@rfsnc.org | Home-care info dept: 919-713-1556 | resourcesforseniors.org/home-care",
            "ask": "Introductory meeting, resource-list discussion, family handoff path",
            "notes": "Most important nonprofit trust anchor on the list.",
        },
        {
            "name": "Meals on Wheels of Wake County",
            "town": "Raleigh",
            "type": "Nonprofit / referral",
            "first_move": "Call first Wednesday morning, then visit",
            "contact": "1001 Blair Dr, Suite 100, Raleigh, NC 27603 | 919-833-1749 | wakemow.org/contact-us",
            "ask": "Intro meeting, resource handoff, possible site-level visibility",
            "notes": "Strong family trust and broad senior touchpoints.",
        },
    ],
    "Priority 2: Community And Caregiver-Support Sites": [
        {
            "name": "Ruth Sheets Adult Care Center",
            "town": "Raleigh",
            "type": "Adult day / caregiver support",
            "first_move": "Call first Monday morning, then visit",
            "contact": "228 W Edenton St, Raleigh, NC 27603 | 919-832-7227 | Matt Frazier | mfrazier@esumc.org | esumc.org/ruth-sheets-adult-care-center",
            "ask": "Intro meeting, support-group mention, flyer permission",
            "notes": "Very strong fit for families already seeking caregiver support.",
        },
        {
            "name": "John M. Brown Community Center",
            "town": "Apex",
            "type": "Community center",
            "first_move": "Call first Tuesday morning, then visit",
            "contact": "53 Hunter St, Apex, NC 27502 | 919-249-3402 | apexnc.org/453/John-M-Brown-Community-Center",
            "ask": "Flyer approval, 55+ programming lead intro",
            "notes": "Good second Apex stop after the senior center.",
        },
        {
            "name": "Bond Park Community Center",
            "town": "Cary",
            "type": "Community center",
            "first_move": "Call first Tuesday morning, then visit",
            "contact": "150 Metro Park Dr, Cary, NC 27513 | 919-462-3970 | carync.gov/recreation-enjoyment/facilities/bond-park-community-center",
            "ask": "Flyer approval, event-table opportunity",
            "notes": "Useful for broad family traffic and civic visibility.",
        },
        {
            "name": "Herbert C. Young Community Center",
            "town": "Cary",
            "type": "Community center",
            "first_move": "Call first Tuesday morning, then visit",
            "contact": "101 Wilkinson Ave, Cary, NC 27513 | 919-460-4965 | carync.gov/recreation-enjoyment/facilities/herbert-c-young-community-center",
            "ask": "Flyer approval, event-table lead",
            "notes": "Good family-facing civic location in Cary.",
        },
        {
            "name": "Laurel Hills Community Center",
            "town": "Raleigh",
            "type": "Community center",
            "first_move": "Call first Monday morning, then visit",
            "contact": "3808 Edwards Mill Rd, Raleigh, NC 27612 | 919-996-2383 | Director Michele Cordaro | raleighnc.gov/parks-and-recreation/places/laurel-hills-park",
            "ask": "Flyer approval, active-adult or family program intro",
            "notes": "Use as a Raleigh family-awareness channel.",
        },
        {
            "name": "John Chavis Memorial Park Community Center",
            "town": "Raleigh",
            "type": "Community center",
            "first_move": "Call first Monday morning, then visit",
            "contact": "505 Martin Luther King Jr Blvd, Raleigh, NC 27601 | 919-996-6590 | raleighnc.gov/parks-and-recreation/places/john-chavis-memorial-park",
            "ask": "Flyer approval, event/program intro, rental or table path",
            "notes": "Good downtown visibility and community reach.",
        },
    ],
    "Priority 3: Church And Network Multipliers": [
        {
            "name": "St. Michael the Archangel Senior Socials",
            "town": "Cary",
            "type": "Church senior ministry",
            "first_move": "Call first Tuesday morning, then visit parish office",
            "contact": "804 High House Rd, Cary, NC 27513 | 919-468-6100 | pastoralcare@stmcary.org | stmichaelcary.org/directory/senior-socials",
            "ask": "Bulletin, pastoral-care referral path, flyer approval",
            "notes": "Strong fit because the ministry already serves seniors 60+.",
        },
        {
            "name": "Cary First Baptist Happy Hearts",
            "town": "Cary",
            "type": "Church senior ministry",
            "first_move": "Call first Tuesday morning, then visit church office",
            "contact": "218 S Academy St, Cary, NC 27511 | 919-467-6356 | office@caryfbc.org | caryfirst.com/happyhearts",
            "ask": "Bulletin, event intro, short caregiver talk",
            "notes": "Official ministry page makes this a good outreach target.",
        },
        {
            "name": "Bay Leaf Baptist Seniors",
            "town": "Raleigh",
            "type": "Church senior ministry",
            "first_move": "Call first Wednesday morning, then visit if open",
            "contact": "12200 Bayleaf Church Rd, Raleigh, NC 27614 | 919-847-4477 | info@bayleaf.org | bayleaf.org/ministries/seniors",
            "ask": "Senior newsletter, flyer approval, ministry intro",
            "notes": "Use the office line first; get the direct ministry lead from staff if possible.",
        },
        {
            "name": "Crossroads Fellowship Seniors",
            "town": "Raleigh",
            "type": "Church senior ministry",
            "first_move": "Call first Wednesday morning, then visit if open",
            "contact": "2721 E Millbrook Rd, Raleigh, NC 27604 | 919-981-0222 | info@crossroads.org | crossroads.org/seniors",
            "ask": "Seniors ministry intro, care ministry intro, bulletin path",
            "notes": "Large church with enough scale to create repeat referrals.",
        },
        {
            "name": "Raleigh Engagement Network",
            "town": "Raleigh",
            "type": "Civic network",
            "first_move": "Call or email after Monday senior-center block",
            "contact": "Lance Shinholser | 919-996-5689 | lance.shinholser@raleighnc.gov | raleighnc.gov/community/services/raleigh-engagement-network",
            "ask": "Explore flyers, mailers, community-center access, and broader visibility",
            "notes": "Official page says network users can access printed flyers/mailers and community-center support.",
        },
        {
            "name": "ONE Wake",
            "town": "Wake County network",
            "type": "Faith-community network",
            "first_move": "Call Thursday morning, request meeting",
            "contact": "919-328-3966 | onewake.org",
            "ask": "Intro to clergy/institution network and caregiver-support fit",
            "notes": "Useful multiplier if they open doors to congregations.",
        },
        {
            "name": "North Carolina Baptist Aging Ministry",
            "town": "Statewide / useful for Wake churches",
            "type": "Faith-based aging network",
            "first_move": "Call Thursday morning, request intro conversation",
            "contact": "877-506-2226 | ncbam.org",
            "ask": "Ask about church-facing caregiver resource alignment and Wake County church pathways",
            "notes": "Official site provides church workshops and caregiver resources.",
        },
    ],
}


def set_cell_text(cell, text):
    cell.text = text
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.size = Pt(9)


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    if level:
        p.paragraph_format.left_indent = Inches(0.25 * level)
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    return p


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    return p


doc = Document()
section = doc.sections[0]
section.left_margin = Inches(0.55)
section.right_margin = Inches(0.55)
section.top_margin = Inches(0.55)
section.bottom_margin = Inches(0.55)

style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(10.5)

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run("HomeCare Wake County Offline GTM Field Manual")
r.bold = True
r.font.size = Pt(18)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = subtitle.add_run("Execution week: April 6-10, 2026")
r.italic = True
r.font.size = Pt(11)

doc.add_paragraph("")

h = doc.add_paragraph()
h.add_run("Purpose").bold = True
add_bullet(doc, "Give the salesperson one clear field document for next week.")
add_bullet(doc, "Focus on offline, trust-based distribution in Wake County.")
add_bullet(doc, "Prioritize senior centers, caregiver-serving organizations, churches, and community multipliers.")

h = doc.add_paragraph()
h.add_run("Data Verification").bold = True
add_bullet(doc, "Target data was checked against official public-facing pages on April 4, 2026.")
add_bullet(doc, "Where a direct named contact was not clearly shown on an official page, this manual uses the main public line or official website instead of guessing.")
add_bullet(doc, "For some churches, use the office line or pastoral-care inbox first, then capture the direct ministry lead during outreach.")

h = doc.add_paragraph()
h.add_run("What Success Looks Like This Week").bold = True
for item in [
    "25 physical stops",
    "40 outbound calls",
    "15 real conversations with staff or decision-makers",
    "10 approved flyer placements",
    "5 follow-up meetings booked",
    "4 talks, table opportunities, or ministry/event placements",
    "3 newsletter or bulletin inclusions",
]:
    add_bullet(doc, item)

h = doc.add_paragraph()
h.add_run("Core Positioning").bold = True
p = doc.add_paragraph()
p.add_run(
    '"HomeCare helps Wake County families quickly find non-medical in-home support for an aging parent. '
    "We are especially useful after discharge, when families need daytime help, or when adult children are coordinating care from work or from out of town.\""
)

h = doc.add_paragraph()
h.add_run("Rules For The Salesperson").bold = True
for item in [
    'Do not start with "Can I leave flyers?" Start with "Who handles family resources or community materials here?"',
    'Do not pitch clinical care. Always say "non-medical home care support."',
    "Do not overexplain the product. Keep the first explanation short and local.",
    "Always ask for the next higher-value step: flyer approval, bulletin/newsletter, event table, short talk, or a direct introduction.",
    "Every stop must end with a next step and a follow-up date.",
]:
    add_number(doc, item)

h = doc.add_paragraph()
h.add_run("Daily Schedule").bold = True

days = [
    (
        "Monday, April 6, 2026",
        "Raleigh senior and active-adult centers plus core referral organizations",
        [
            "8:00-9:00 a.m.: Call Five Points, Anne Gordon, Laurel Hills, John Chavis, Resources for Seniors, and Ruth Sheets.",
            "9:30 a.m.-12:00 p.m.: Visit Five Points, Anne Gordon, and Resources for Seniors.",
            "1:00-4:30 p.m.: Visit Ruth Sheets, Laurel Hills, and John Chavis.",
            "4:30-5:30 p.m.: Log outcomes, send follow-up emails, and confirm Tuesday appointments.",
        ],
    ),
    (
        "Tuesday, April 7, 2026",
        "Cary and Apex senior-focused outreach",
        [
            "8:00-9:00 a.m.: Call Cary Senior Center, Bond Park, Herbert C. Young, St. Michael Cary, Cary First Baptist, Apex Senior Center, and John M. Brown.",
            "9:30 a.m.-12:30 p.m.: Visit Cary Senior Center, Bond Park, and Herbert C. Young.",
            "1:30-5:00 p.m.: Visit St. Michael Cary, Cary First Baptist, Apex Senior Center, and John M. Brown.",
        ],
    ),
    (
        "Wednesday, April 8, 2026",
        "Garner and Wake Forest, plus nonprofit and church follow-up",
        [
            "8:00-9:00 a.m.: Call Garner Senior Center, Wake Forest Center for Active Aging, Meals on Wheels, Bay Leaf Baptist, and Crossroads.",
            "9:30 a.m.-12:00 p.m.: Visit Garner Senior Center and Meals on Wheels.",
            "1:30-4:30 p.m.: Visit Wake Forest Center for Active Aging, Bay Leaf Baptist, and Crossroads.",
        ],
    ),
    (
        "Thursday, April 9, 2026",
        "Church and network day",
        [
            "8:00-9:00 a.m.: Call remaining church leads, ONE Wake, NCBAM, and all warm follow-ups.",
            "9:30 a.m.-12:30 p.m.: Visit warm church offices or attend booked conversations.",
            "1:30-5:00 p.m.: Make second visits to best leads and deliver requested materials.",
        ],
    ),
    (
        "Friday, April 10, 2026",
        "Follow-up and conversion day",
        [
            "8:00-10:00 a.m.: Call every warm lead and lock next-week follow-up.",
            "10:00 a.m.-12:00 p.m.: Revisit the highest-probability institutions.",
            "1:00-3:30 p.m.: Final visit block for warm accounts only.",
            "3:30-5:30 p.m.: Complete weekly report and rank the top 10 warm accounts.",
        ],
    ),
]

for day, theme, bullets in days:
    p = doc.add_paragraph()
    run = p.add_run(day)
    run.bold = True
    run.font.size = Pt(11.5)
    p = doc.add_paragraph()
    p.add_run(f"Theme: {theme}")
    for bullet in bullets:
        add_bullet(doc, bullet)

h = doc.add_paragraph()
h.add_run("What To Say").bold = True

script_sections = [
    (
        "Front desk / first contact",
        "Hi, I am with HomeCare. We help Wake County families find non-medical in-home support for an aging parent. Who handles family resource information or community materials here?",
    ),
    (
        "Senior center director or program lead",
        "We are building awareness in Wake County for families who suddenly need non-medical support at home for an aging parent. We already have caregiver supply locally, and we are trying to become a practical local resource rather than just another ad. Would you be open to letting us leave materials, submit a newsletter blurb, or discuss a short informational talk?",
    ),
    (
        "Church office",
        "Hi, I am reaching out from HomeCare. We help Wake County families find non-medical support for an aging parent at home. I am looking for the person who handles member care, pastoral care, senior adult ministry, or church bulletin announcements. Who would that be?",
    ),
    (
        "Pastor or pastoral-care lead",
        "A lot of adult children in churches are quietly trying to coordinate help for an aging parent, especially after a health event or when work makes daily support hard. We would love to be a practical local resource. Would it be useful to leave materials, submit a bulletin blurb, or offer a short caregiver-focused session?",
    ),
    (
        "Nonprofit / referral organization",
        "We know you already support older adults and caregivers in Wake County. We are not asking for a hard referral right away. We would simply value an introductory conversation to explain where HomeCare may fit as a private-pay non-medical option when families need help quickly.",
    ),
]

for heading, text in script_sections:
    p = doc.add_paragraph()
    p.add_run(heading).bold = True
    doc.add_paragraph(text)

h = doc.add_paragraph()
h.add_run("Objection Handling").bold = True
for item in [
    'If they say "We do not allow flyers," respond: "No problem. Do you ever share local resources through a newsletter, email, bulletin, or event table?"',
    'If they say "We cannot endorse outside businesses," respond: "Understood. We are simply asking whether HomeCare can be listed as a local resource families may choose to explore, not as an endorsed provider."',
    'If they say "Send me something by email," respond: "Absolutely. What is the best email, and when should I follow up?"',
]:
    add_bullet(doc, item)

h = doc.add_paragraph()
h.add_run("Target Directory").bold = True
doc.add_paragraph("Use this section in the field. First move tells the rep exactly what to do first.")

for section_name, rows in TARGETS.items():
    doc.add_page_break()
    p = doc.add_paragraph()
    p.add_run(section_name).bold = True
    p.runs[0].font.size = Pt(12)

    table = doc.add_table(rows=1, cols=7)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    headers = ["Organization", "Town", "Type", "First Move", "Contact", "Primary Ask", "Notes"]
    for i, header in enumerate(headers):
        set_cell_text(hdr[i], header)
        shade_cell(hdr[i], "D9EAF7")
        for run in hdr[i].paragraphs[0].runs:
            run.bold = True

    for row in rows:
        cells = table.add_row().cells
        set_cell_text(cells[0], row["name"])
        set_cell_text(cells[1], row["town"])
        set_cell_text(cells[2], row["type"])
        set_cell_text(cells[3], row["first_move"])
        set_cell_text(cells[4], row["contact"])
        set_cell_text(cells[5], row["ask"])
        set_cell_text(cells[6], row["notes"])

h = doc.add_paragraph()
h.add_run("Same-Day Follow-Up Email").bold = True
doc.add_paragraph(
    "Subject: Local Wake County family care resource\n\n"
    "Hello [Name],\n\n"
    "Thank you for speaking with me today. HomeCare helps Wake County families find non-medical in-home support for an aging parent, especially after discharge, during daytime care gaps, or when adult children are coordinating care from work or from out of town.\n\n"
    "As discussed, I am sending over a short overview, a family flyer, and a short blurb you can use in a bulletin or newsletter if helpful. We would also be glad to offer a short educational session for families or caregivers on how to find non-medical support at home in Wake County.\n\n"
    "Thank you again,\n[Name]\n[Phone]\n[Email]"
)

h = doc.add_paragraph()
h.add_run("End-Of-Day Reporting").bold = True
for item in [
    "Stops completed",
    "Calls completed",
    "Decision-makers reached",
    "Flyer placements approved",
    "Meetings booked",
    "Bulletins/newsletters submitted",
    "Best lead of the day",
    "Top objection heard",
    "Must-follow-up tomorrow",
]:
    add_bullet(doc, item)

doc.save(OUTPUT_PATH)
print(OUTPUT_PATH)
